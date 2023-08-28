import openpyxl
from openpyxl import Workbook
from datetime import datetime
# from pymysql import *
# import pandas.io.sql as sql
import pandas as pd
import os
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT

class LoadStrategy:

    stored_values = ''

    def __init__(self, filename, input_folder_name, output_folder_name, run_time):
        self.input_folder_name = input_folder_name
        self.filename = filename
        self.output_folder_name = output_folder_name
        self.run_time = run_time

    def read_file(self,status,conn):
        print("Load Strategy")
        wb = openpyxl.load_workbook('%s\\Input Files\\%s' %(self.input_folder_name,self.filename))
        lg_sheet = wb['UserInput']
        max_cl = lg_sheet.max_column
        for i in range(5,lg_sheet.max_row+1):
            data_dic = {}
            for j in range(1,max_cl+1):
                data_dic[lg_sheet.cell(row=1,column=j).value] = str(lg_sheet.cell(row=i,column=j).value)
            for key,value in data_dic.items():
                if '_DT' in key or '_COL' in key:
                    if value is not None or str(value).replace(' ','') == '':
                        data_dic[key] =  str(value)
                    else:
                        data_dic[key] = ' '
            print(data_dic)
            self.performLoadStrategy(wb, data_dic,data_dic['Load Strategy'],status,conn)


    def performLoadStrategy(self,wb,data,lg_type,status,conn):

        folder = self.output_folder_name + '\\Output Files\\' + 'LoadStrategy\\%s' % status
        if not os.path.exists(folder):
            os.makedirs(folder)
        html_data = """<!DOCTYPE html>
                <html>
                <head>
                <meta name="viewport" content="width=device-width, initial-scale=1">
                <style>
                table {
                  border-collapse: collapse;
                  border-spacing: 0;
                  width: 50%;
                  height : 70px;
                  overflow-x : auto;
                  border: 2px solid #ddd;
                  font-size: 15px;
                }
                th, td {
                  text-align: center;
                  height:2px;
                  width:auto;
                  border: black solid 3px;
                  padding: 2px;
                }
                p{
                  width: 100%;
                  height : auto;
                  border: 3px solid #ddd;
                  overflow: auto;
                   }
                #tabledata{height:200px;width:50%;overflow:auto;}
                tr:nth-child(even){background-color: #f2f2f2 ; width:10px;height:50px}
                </style>
                </head>
                <body>
                <h2 style="text-align:center"><<Table_Name>></h2>
                <hr/>
                """
        middle = """<b><h2>TableName: <TABLENAME> - Layer:&nbsp<i><LAYER></i><h2></b><hr/>
                <h3>Load Strategy : <LOADTYPE></h3>
                <div style = 'border:3px solid black;padding: 10px;background-color:white'><br/>
                <h3>Status: <i><STATUS</i></h3>
                <i>QUERY:</i>
                <p><<QUERY>></p>
                <hr/><hr/>
                <i>OUTPUT:&nbsp <RESULT></i>
                <hr/>
                <div id = 'tabledata'>
                  <<QUERIED_DATAFRAME_OUTPUT>>
                </div>
                </div>"""

        footer = """</body>
                        </html>"""
        document = Document()
        document.add_heading(data['Source Table Name']+' : '+data['Layer']+' : '+data['Load Strategy'], 0)
        if lg_type.upper() == 'TRUNCATE':
            sheet = wb['Truncate']
        elif lg_type.upper() == 'APPEND':
            sheet = wb['Append']
        elif lg_type.upper() == 'SCD1':
            sheet = wb['SCD1']
        # print(data)
        temp_middle = ''
        val1 = ''
        val2 = ''
        stored_values = ''
        for i in range(2,sheet.max_row+1):
            type = sheet.cell(row=i, column=1).value
            job_status = sheet.cell(row=i, column=2).value
            query = sheet.cell(row=i, column=3).value
            if status.upper()  in job_status.upper():
                if type.upper() == 'QUERY':
                    # print(query)
                    temp_path = self.output_folder_name + '\\Output Files\\' + 'LoadStrategy\\temp.txt'
                    output = ['','']
                    if os.path.exists(temp_path):
                        file1 = open(temp_path,"r+")
                        output = file1.read()
                        output = output.split(',')
                    val1 = output[0]
                    val2 = output[1]
                    query1 = query.replace('<SRC_TBL>', data['Source Table Name']).replace('<TGT_TBL>', data[
                        'Target Table Name']).replace('<SRC_KEY_COL>', data['SRC_KEYCOLUMN']).replace('<SRC_CRT_DT>',
                                                                                                      data[
                                                                                                          'SRC_CRT_DT']) \
                        .replace('<SRC_UPD_DT>', data['SRC_UPD_DT']).replace('<SRC_ADD_COL>', data['SRC_ADD_COL']) \
                        .replace('<TGT_KEY_COL>', data['TGT_KEYCOLUMN']).replace('<TGT_CRT_DT>',
                                                                                 data['TGT_CRT_DT']).replace(
                        '<TGT_UPD_DT>', data['TGT_UPD_DT']) \
                        .replace('<TGT_UPD_COL_VAL>', data['TGT_UPD_COL_VAL']).replace('<TGT_ADD_COL>',
                                                                                       data['TGT_ADD_COL_VAL']).replace(
                        ', None', '').replace('<VAL1>', val1).replace('<VAL2>', val2)
                    print(query1)
                    # write execute query code
                    db_df = pd.read_csv(
                        r"C:\Users\MadhReddy\Desktop\AwsTestingAccelerator\Input Files\Monthly_Input_Home_RAW.csv",
                         index_col=False)
                    # db_df = sql.read_sql(query1, conn)
                    headers = list(db_df.columns)
                    records = db_df.values.tolist()
                    stored_values = records
                    document.add_heading('Query: %s : ' % (query1), level=1)
                    table = document.add_table(rows=1, cols=len(headers), style='Table Grid')
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    hdr_cells = table.rows[0].cells
                    for i in range(len(headers)):
                        hdr_cells[i].text = headers[i]
                    for record in records:
                        row_cells = table.add_row().cells
                        for i in range(len(record)):
                            row_cells[i].text = str(record[i])
                    temp = db_df.to_html(index=False).replace(',', '').replace('(', '').replace(')', '')
                    temp_middle1 = temp_middle.replace(
                        "<<QUERIED_DATAFRAME_OUTPUT>>",
                        temp).replace("<<QUERY>>",
                                      query).replace(
                        '<LAYER>', data['Layer']).replace('<TABLENAME>', data['Source Table Name']).replace(
                        '<LOADTYPE>', data['Load Strategy'])

                    middle = temp_middle1 + middle
                elif type.upper() == 'JOB':
                    print("Trigger the Job")
                    # write code to run the job
                elif type.upper() == 'STORE':
                    print("Store")
                    val1 = records[0][0]
                    val2 = records[1][0]
                    file1 = open(self.output_folder_name + '\\Output Files\\' + 'LoadStrategy\\temp.txt', "w")
                    file1.write(val1+',')
                    file1.write(val2)
                    file1.close()

                    # stored_values = db_df.values.tolist()['TGT_KEY_COL']
        run_time = str(datetime.now()).replace(':', '').replace(' ', '').replace('.', '')
        # final_report = html_data+middle+footer
        # file = open(folder + '\\LoadStrategy_%s_%s.html' % (data['Source Table Name'], run_time), mode='w', encoding="utf-8")
        # file.write(final_report)
        # file.close()
        document.save(folder + '\\LoadStrategy_%s_%s.docx' % (data['Source Table Name'], run_time))

    #
    # def generatereports(self):
    #     html_data = """<!DOCTYPE html>
    #     <html>
    #     <head>
    #     <meta name="viewport" content="width=device-width, initial-scale=1">
    #     <style>
    #     table {
    #       border-collapse: collapse;
    #       border-spacing: 0;
    #       width: 50%;
    #       height : 70px;
    #       overflow-x : auto;
    #       border: 2px solid #ddd;
    #       font-size: 15px;
    #     }
    #     th, td {
    #       text-align: center;
    #       height:2px;
    #       width:auto;
    #       border: black solid 3px;
    #       padding: 2px;
    #     }
    #     p{
    #       width: 100%;
    #       height : auto;
    #       border: 3px solid #ddd;
    #       overflow: auto;
    #        }
    #     #tabledata{height:200px;width:50%;overflow:auto;}
    #     tr:nth-child(even){background-color: #f2f2f2 ; width:10px;height:50px}
    #     </style>
    #     </head>
    #     <body>
    #     <h2 style="text-align:center"><<Table_Name>></h2>
    #     <hr/>
    #     """
    #     middle = """<h3>TableName: <TABLENAME> - Layer:&nbsp<i><LAYER></i></h3>
    #     <div style = 'border:3px solid black;padding: 10px;background-color:white'>
    #     <i>QUERY:</i>
    #     <p><<QUERY>></p>
    #     <hr/><hr/>
    #     <i>OUTPUT:&nbsp <RESULT></i>
    #     <hr/>
    #     <div id = 'tabledata'>
    #       <<QUERIED_DATAFRAME_OUTPUT>>
    #     </div>
    #     </div>"""
    #
    #     footer = """</body>
    #             </html>"""







# run_time = str(datetime.now()).replace(':','').replace(' ','').replace('.','')
# input_folder_name  = r'C:\Users\MadhReddy\Desktop\Takeda Automation'
# output_folder_nm = r'C:\Users\MadhReddy\Desktop\DataValidations\Queries'
# file_name = r"C:\Users\MadhReddy\Desktop\Takeda Automation\Load Strategy.xlsx"
# lg = LoadStrategy(file_name,input_folder_name,output_folder_nm, run_time)
# lg.read_file('Before')
# lg.read_file('After')