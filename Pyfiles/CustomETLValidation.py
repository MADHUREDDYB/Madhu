import urllib3
from urllib3 import disable_warnings
import openpyxl
from openpyxl import Workbook
from openpyxl.styles import PatternFill, Alignment, Font
import pandas as pd
import os
import datetime
import collections
# import jaydebeapi as jdbc
import json
import jpype
from pymysql import *
import pandas.io.sql as sql
import csv
import glob
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT


class ValidateData:

    def __init__(self, conn, output_folder_nm, run_time):
        self.output_folder_nm = output_folder_nm
        self.conn = conn
        self.run_time = run_time

    def GenerateReport(self):
        folder = self.output_folder_nm + '\\Output Files\\Custom Testing\\ETL Validations'
        if not os.path.exists(folder):
            os.makedirs(folder)
        filename = self.output_folder_nm + '\\Output Files\\Custom Testing\\Generated_ETL_Queries'
        # list_of_files = os.listdir(filename)  # * means all if need specific format then *.csv
        files = os.listdir(filename)
        if not files:
            return '<b>EXECUTION STOPPED! : </b>The  Queries are not present at %s: ' % filename, folder
        files = files[::-1][0]
        filename = filename + '\\' + files
        print(filename)
        wb = openpyxl.load_workbook(filename)
        wb_output = Workbook()
        html_data = """<!DOCTYPE html>
<html>
<head>
<meta name="viewport" content="width=device-width, initial-scale=1">
<style>
table {
  border-collapse: collapse;
  border-spacing: 0;
  width: 50%;
  height : 70px;
  overflow-x : auto;
  border: 2px solid #ddd;
  font-size: 15px;
}
th, td {
  text-align: center;
  height:2px;
  width:auto;
  border: black solid 3px;
  padding: 2px;
}
p{
  width: 100%;
  height : auto;
  border: 3px solid #ddd;
  overflow: auto;
   }
#tabledata{height:200px;width:50%;overflow:auto;}
tr:nth-child(even){background-color: #f2f2f2 ; width:10px;height:50px}
</style>
</head>
<body>
<h2 style="text-align:center"><<Table_Name>></h2>
<hr/>
<h2> Summary :</h2>
<table style="width:100%;border 3px black">
  <tr>
    <th>Count Type</th>
    <th>Count</th>
  </tr>
  <tr>
    <td>Total Test Cases</td>
    <td><<totalcount>></td>
  </tr>
    <tr>
    <td>Test Case(s) Not Executed</td>
    <td><<nonexecount>></td>
  </tr>
  <tr>
    <td>Test Case(s) Executed</td>
    <td><<execount>></td>
  </tr>
  <tr style = 'color:darkgreen'>
    <td>Test Case(s) Passed</td>
    <td><<passedcount>></td>
  </tr>
  <tr style = 'color:Red'>
    <td>Test Case(s) Failed</td>
    <td><<failedcount>></td>
  </tr>
</table>
<hr/>
"""
        middle = """<h3><i><TS_ID></i> - <TC_ID> : <<TC_DESC>></h3>
<div style = 'border:3px solid black;padding: 10px;background-color:white'>
<i>QUERY:</i>
<p><<QUERY>></p>
<hr/><hr/>
<i>OUTPUT:&nbsp <RESULT></i>
<hr/>
<div id = 'tabledata'>
  <<QUERIED_DATAFRAME_OUTPUT>>
</div>
</div>"""

        footer = """</body>
        </html>"""
        tb_name = ''
        summary = ''
        final_count = [0] * 5
        for sheet in wb.sheetnames:
            document = Document()
            total = passed = failed = executed = notexecuted = 0
            queries_sheet = wb[sheet]
            report_data = ""
            tb_name = queries_sheet.cell(row=2, column=3).value
            document.add_heading(tb_name, 0)
            document.add_heading('Summary', level=1)
            table1 = document.add_table(rows=2, cols=2, style='Table Grid')
            hdr_cells1 = table1.rows[0].cells
            hdr_cells1[0].text = 'Count Type'
            hdr_cells1[1].text = 'Count'
            row_cells1 = table1.add_row().cells
            row_cells1[0].text = 'Total Test Cases'
            row_cells2 = table1.add_row().cells
            row_cells2[0].text = 'Test Case(s) Executed'
            row_cells3 = table1.add_row().cells
            row_cells3[0].text = 'Test Case(s) Not Executed'
            row_cells4 = table1.add_row().cells
            row_cells4[0].text = 'Test Case(s) Passed'
            row_cells5 = table1.add_row().cells
            row_cells5[0].text = 'Test Case(s) Failed'
            temp = tb_name.split('.')
            ws1 = wb_output.create_sheet(sheet)
            ws1.freeze_panes = 'B2'
            ws1.sheet_view.zoomScale = 70
            ws1.auto_filter.ref = 'A1:M1'
            ws1.column_dimensions['A'].width = 10
            ws1.column_dimensions['B'].width = 10
            ws1.column_dimensions['C'].width = 15
            ws1.column_dimensions['D'].width = 15
            ws1.column_dimensions['E'].width = 10
            ws1.column_dimensions['F'].width = 20
            ws1.column_dimensions['G'].width = 30
            ws1.column_dimensions['H'].width = 50
            ws1.column_dimensions['I'].width = 10
            ws1.column_dimensions['J'].width = 10
            ws1.column_dimensions['K'].width = 10
            ws1.column_dimensions['L'].width = 10
            ws1.column_dimensions['M'].width = 10
            headers = ['Test Scenario ID', 'Test Case ID', 'Table Name', 'Column Name', 'Level', 'Test Case Name',
                       'Test Case Desc', 'Query', 'Exe Flag', 'Exe Status', 'Execpted Value', 'Actual Value', 'Result']
            ws1.append(headers)
            for m in range(1, len(headers) + 1):
                ws1.cell(row=1, column=m).fill = PatternFill(start_color="575757", end_color="575757",
                                                             fill_type="solid")
                ws1.cell(row=1, column=m).font = Font(color='6FF242')
            total = queries_sheet.max_row - 1
            for rownum in range(2, queries_sheet.max_row + 1):
                ts_id = queries_sheet.cell(row=rownum, column=1).value
                tc_id = queries_sheet.cell(row=rownum, column=2).value
                column_nm = level = queries_sheet.cell(row=rownum, column=4).value
                level = queries_sheet.cell(row=rownum, column=5).value
                tc_nm = str(queries_sheet.cell(row=rownum, column=6).value)
                tc_desc = queries_sheet.cell(row=rownum, column=7).value
                query = queries_sheet.cell(row=rownum, column=8).value
                exe_flag = queries_sheet.cell(row=rownum, column=9).value
                exp_value = queries_sheet.cell(row=rownum, column=11).value
                res = ''
                temp = ''
                temp2 = ''
                ws1.cell(row=rownum, column=1).value = ts_id
                ws1.cell(row=rownum, column=2).value = tc_id
                ws1.cell(row=rownum, column=3).value = tb_name
                if level == 'Table Level':
                    ws1.cell(row=rownum, column=4).value = 'All Columns'
                    ws1.cell(row=rownum, column=5).value = "Table Level"
                else:
                    ws1.cell(row=rownum, column=4).value = column_nm
                    ws1.cell(row=rownum, column=5).value = "Column Level"
                ws1.cell(row=rownum, column=6).value = tc_nm
                ws1.cell(row=rownum, column=7).value = tc_desc
                ws1.cell(row=rownum, column=8).value = query
                ws1.cell(row=rownum, column=8).alignment = Alignment(wrap_text=True)
                ws1.cell(row=rownum, column=9).value = exe_flag
                ws1.cell(row=rownum, column=11).value = exp_value
                records, headers = [], []
                if str(exe_flag).upper() == 'TRUE':
                    executed += 1
                    try:
                        query_output = sql.read_sql(query, self.conn)
                        d = query_output.to_string(index=False, header=False)
                        headers = list(query_output.columns)
                        records = query_output.values.tolist()
                        ws1.cell(row=rownum, column=12).value = d
                        if tc_nm.upper().replace(" ", '') == 'TABLECOUNTCHECK':
                            d = query_output['COUNT'].tolist()
                            if d[0] == d[1]:
                                res = 'Pass'
                            else:
                                res = 'Fail'
                        elif 'MINUS' in tc_nm.upper() or tc_nm == 'Dupliacte Check' or 'Job statistics' in tc_nm:
                            if str(d) == '0':
                                res = "Pass"
                            else:
                                res = 'Fail'
                        elif tc_nm.upper().replace(" ", '') == 'KEYCOLUMNCHECK':
                            d = query_output['cnt'].tolist()
                            if len(set(d)) == 1:
                                res = "Pass"
                            else:
                                res = 'Pass'
                        else:
                            res = 'Pass'

                        if res == 'Pass':
                            passed += 1
                            ws1.cell(row=rownum, column=13).value = 'Pass'
                        else:
                            failed += 1
                            res = 'Fail'
                            ws1.cell(row=rownum, column=13).value = 'Fail'
                        ws1.cell(row=rownum, column=12).value = query_output.to_string(index=False,
                                                                                       header=False)  # edited
                        ws1.cell(row=rownum, column=10).value = "Executed"
                        temp = middle
                        temp2 = query_output.to_html(index=False).replace(',', '').replace('(', '').replace(')', '')

                    except Exception as E:
                        res = 'Failed Due to Database Error : Error Description'
                        # print(E)
                        headers = ['Failure Summary']
                        records = [[E]]
                        failed += 1
                        ws1.cell(row=rownum, column=10).value = "Not Executed"
                        ws1.cell(row=rownum, column=12).value = str(E)
                        temp = middle
                        temp2 = str(E)
                        ##########adding doc Code here
                    document.add_heading('%s_%s  - Result: %s ' % (ts_id, tc_id, res), level=1)
                    document.add_paragraph(tc_desc, style='Intense Quote')
                    table = document.add_table(rows=len(records), cols=len(headers), style='Table Grid')
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    hdr_cells = table.rows[0].cells
                    print(headers)
                    print(records)
                    print(len(records))
                    for i in range(len(headers)):
                        hdr_cells[i].text = headers[i]
                    for record in records:
                        row_cells = table.add_row().cells
                        for i in range(len(record)):
                            row_cells[i].text = str(record[i])

                        # document.add_page_break()

                    report_data = report_data + temp.replace("<<TC_DESC>>", tc_desc).replace(
                        "<<QUERIED_DATAFRAME_OUTPUT>>",
                        temp2).replace("<<QUERY>>",
                                       query).replace(
                        '<TS_ID>', ts_id).replace('<TC_ID>', tc_id).replace('<RESULT>', res)
                else:
                    notexecuted += 1
            final_count[0] = final_count[0] + total
            final_count[1] = final_count[1] + executed
            final_count[2] = final_count[2] + notexecuted
            final_count[3] = final_count[3] + passed
            final_count[4] = final_count[4] + failed
            row_cells1[1].text = str(total)
            row_cells2[1].text = str(executed)
            row_cells3[1].text = str(notexecuted)
            row_cells4[1].text = str(passed)
            row_cells5[1].text = str(failed)
            final_report = html_data.replace('<<Table_Name>>', tb_name).replace('<<totalcount>>', str(total)).replace(
                '<<execount>>', str(executed)).replace('<<nonexecount>>', str(notexecuted)).replace('<<passedcount>>',
                                                                                                    str(
                                                                                                        passed)).replace(
                '<<failedcount>>', str(failed)) + report_data.replace('NaN', '') + footer
            document.save(folder + '\\ETLValidations_%s_%s.docx' % (tb_name, self.run_time))
            file = open(folder + '\\ETLValidations_%s_%s.html' % (tb_name, self.run_time), mode='w', encoding="utf-8")
            file.write(final_report)
            file.close()
        summary = "<b>Summary:Total : %s ; Executed : %s ; NotExecuted : %s ; Passed : %s ; Failed : %s </b><BR> " % (
            final_count[0],
            final_count[1],
            final_count[2],
            final_count[3],
            final_count[4])
        wb_output.remove(wb_output['Sheet'])
        wb_output.save(folder + '\\Custom_ETLValidations_%s.xlsx' % (self.run_time))
        return summary, folder


if __name__ == "__GenerateReport__":
    GenerateReport()
